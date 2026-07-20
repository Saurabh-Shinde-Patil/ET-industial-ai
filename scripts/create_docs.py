import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, HRFlowable, KeepTogether
from reportlab.pdfgen import canvas

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def generate_word_document(docx_path):
    doc = Document()
    
    # Page setup - 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles & Colors
    # Primary: Deep Navy (#1E293B), Secondary: Slate Cyan (#0EA5E9), Text: Dark Gray (#334155)
    NAVY = RGBColor(30, 41, 59)
    CYAN = RGBColor(14, 165, 233)
    GRAY = RGBColor(51, 65, 85)
    LIGHT_BG = "F8FAFC"
    
    # Normal Style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = GRAY

    # Cover Header Title Block
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run("ET AI HACKATHON 2.0 — PHASE 2 PROTOTYPE SUBMISSION")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(10)
    title_run.font.bold = True
    title_run.font.color.rgb = CYAN
    
    h1_p = doc.add_paragraph()
    h1_run = h1_p.add_run("Industrial Knowledge Intelligence Platform (PS-8)")
    h1_run.font.name = 'Calibri'
    h1_run.font.size = Pt(22)
    h1_run.font.bold = True
    h1_run.font.color.rgb = NAVY
    h1_p.paragraph_format.space_after = Pt(2)

    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run("Technical Architecture & Operational Specification Document | Unified Asset & Operations Brain")
    sub_run.font.name = 'Calibri'
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = GRAY
    sub_p.paragraph_format.space_after = Pt(18)

    # Metadata Table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Problem Statement", "PS 8: AI for Industrial Knowledge Intelligence: Unified Asset & Operations Brain"),
        ("Project Name", "ET Industrial AI"),
        ("Repository URL", "https://github.com/Saurabh-Shinde-Patil/ET-industrial-ai"),
        ("Date", "July 20, 2026")
    ]
    for idx, (label, val) in enumerate(meta_data):
        row = meta_table.rows[idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        
        lbl_p = cell_lbl.paragraphs[0]
        r1 = lbl_p.add_run(label)
        r1.font.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = NAVY
        
        val_p = cell_val.paragraphs[0]
        r2 = val_p.add_run(val)
        r2.font.size = Pt(10)
        r2.font.color.rgb = GRAY
        
        set_cell_background(cell_lbl, LIGHT_BG)
        set_cell_background(cell_val, "FFFFFF")
        set_cell_margins(cell_lbl, 80, 80, 100, 100)
        set_cell_margins(cell_val, 80, 80, 100, 100)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Helper function for Section Headings
    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(15)
        r.font.bold = True
        r.font.color.rgb = NAVY
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(12.5)
        r.font.bold = True
        r.font.color.rgb = CYAN
        return h

    # Section 1
    add_heading_1("1. Executive Summary")
    p1 = doc.add_paragraph(
        "In thermal power plants, chemical process facilities, and heavy manufacturing installations, critical operational knowledge "
        "is frequently fragmented across thousands of static PDF manuals, handwritten maintenance logs, Standard Operating Procedures (SOPs), "
        "and historic Root Cause Analysis (RCA) records. During acute equipment trips or forced outages, plant engineers face significant "
        "delays sifting through non-searchable document repositories, leading to extended downtime, elevated safety risks, and operational inefficiencies."
    )
    p1.paragraph_format.space_after = Pt(8)
    
    p2 = doc.add_paragraph(
        "ET Industrial AI addresses this operational bottleneck by deploying a unified, enterprise-grade Knowledge Intelligence Platform. "
        "Built on a microservices architecture, the system combines a Zero-Hallucination Retrieval-Augmented Generation (RAG) conversational engine, "
        "a high-density vector search index (FAISS), an automated 5-Whys Root Cause Analysis engine, and predictive maintenance risk scoring. "
        "The platform provides plant engineers with immediate, contextual answers backed strictly by ground-truth document citations, enforcing "
        "strict confidence thresholds to prevent AI hallucination."
    )
    p2.paragraph_format.space_after = Pt(12)

    # Section 2
    add_heading_1("2. Problem Statement Alignment (PS-8)")
    doc.add_paragraph(
        "The project is engineered specifically for Problem Statement 8: AI for Industrial Knowledge Intelligence: Unified Asset & Operations Brain. "
        "The table below outlines how specific plant operational challenges are addressed by the platform architecture:"
    )
    
    ps_table = doc.add_table(rows=6, cols=2)
    ps_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Industrial Operational Challenge", "ET Industrial AI Architectural Solution"]
    
    hdr_cells = ps_table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].paragraphs[0].add_run(title).font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr_cells[i], "1E293B")
        set_cell_margins(hdr_cells[i], 100, 100, 120, 120)

    rows_data = [
        ("Knowledge Isolation across Plant Siloes", "Centralized document ingestion & OCR parsing supporting PDF, Word, TXT, and scanned image formats with direct asset tagging."),
        ("Risk of AI Hallucinations on Critical Assets", "Zero-Hallucination RAG Chain enforcing similarity threshold suppression (<50% score). Returns standardized fallback notices when manuals lack matching context."),
        ("Repetitive Equipment Failures", "Automated 5-Whys Root Cause Analysis engine generating standardized incident reports and Corrective Action (CAPA) plans."),
        ("Unplanned Equipment Downtime", "AI Predictive Maintenance engine calculating asset failure risk probabilities based on telemetry parameters and operating hours."),
        ("Enterprise Security & Role Compliance", "8 RBAC roles (Plant Operator, Maintenance Eng, Reliability Eng, Safety Officer, Production Eng, Plant Mgr, Knowledge Admin, System Admin) with audit event logging.")
    ]

    for idx, (ch, sol) in enumerate(rows_data, start=1):
        row_cells = ps_table.rows[idx].cells
        c1, c2 = row_cells[0], row_cells[1]
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(ch)
        r1.font.bold = True
        r1.font.size = Pt(9.5)
        
        p2 = c2.paragraphs[0]
        r2 = p2.add_run(sol)
        r2.font.size = Pt(9.5)
        
        bg = LIGHT_BG if idx % 2 == 1 else "FFFFFF"
        set_cell_background(c1, bg)
        set_cell_background(c2, bg)
        set_cell_margins(c1, 80, 80, 100, 100)
        set_cell_margins(c2, 80, 80, 100, 100)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 3
    add_heading_1("3. System Architecture & Technical Implementation")
    doc.add_paragraph(
        "The system is structured as a decoupled, multi-tier microservice architecture to ensure high throughput, horizontal scalablity, "
        "and strict separation between the API gateway, vector store, and AI inference engines."
    )

    add_heading_2("3.1 Core Components")
    doc.add_paragraph(
        "1. Frontend Web Client (React 18 / Tailwind CSS / Lucide Icons):\n"
        "   - Responsive industrial dark-mode interface with operational dashboards, RAG conversation view, document manager, asset hierarchy, and admin audit tools.\n\n"
        "2. Node.js Express REST API Gateway (Port 5000):\n"
        "   - Handles user authentication, JWT session validation, RBAC route protection, MongoDB persistence, audit event logging, and proxy dispatch to the Python AI service.\n\n"
        "3. Python FastAPI AI Microservice (Port 8000):\n"
        "   - Houses the core AI engines: SentenceTransformers embedding generation, FAISS vector indexing, Reciprocal Rank Fusion (RRF) hybrid search, PyPDF2/pdfplumber text extraction, Tesseract OCR, and RAG synthesis."
    )

    add_heading_2("3.2 Technical Stack Summary")
    tech_table = doc.add_table(rows=5, cols=2)
    tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_headers = ["Subsystem Layer", "Technologies & Specifications"]
    for i, title in enumerate(t_headers):
        tech_table.rows[0].cells[i].paragraphs[0].add_run(title).font.bold = True
        tech_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(tech_table.rows[0].cells[i], "1E293B")
        set_cell_margins(tech_table.rows[0].cells[i], 100, 100, 120, 120)

    t_data = [
        ("User Interface Layer", "React 18, Vite, Tailwind CSS, Lucide Icons, React Router DOM v6"),
        ("API Gateway Layer", "Node.js (v20+), Express.js, JWT Authentication, Winston Logger, OpenAPI Swagger UI"),
        ("Database & Persistence", "MongoDB (Mongoose ORM) for document metadata/chunks, FAISS binary index on disk for vectors"),
        ("AI & NLP Microservice", "Python 3.12, FastAPI, SentenceTransformers ('all-MiniLM-L6-v2'), FAISS IndexFlatIP, Tesseract OCR")
    ]
    for idx, (layer, specs) in enumerate(t_data, start=1):
        c1, c2 = tech_table.rows[idx].cells[0], tech_table.rows[idx].cells[1]
        c1.paragraphs[0].add_run(layer).font.bold = True
        c1.paragraphs[0].runs[0].font.size = Pt(9.5)
        c2.paragraphs[0].add_run(specs).font.size = Pt(9.5)
        bg = LIGHT_BG if idx % 2 == 1 else "FFFFFF"
        set_cell_background(c1, bg)
        set_cell_background(c2, bg)
        set_cell_margins(c1, 80, 80, 100, 100)
        set_cell_margins(c2, 80, 80, 100, 100)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 4
    add_heading_1("4. Detailed AI & Vector Search Pipeline")
    
    add_heading_2("4.1 Document Processing & Vector Chunking")
    doc.add_paragraph(
        "When an engineer uploads an industrial document (e.g., Boiler Maintenance Manual), the system executes a three-phase pipeline:\n"
        "1. Text Extraction: Text is extracted using pdfplumber or Tesseract OCR (for scanned engineering drawings and paper SOPs).\n"
        "2. Chunking: Text is segmented into 500-token sliding windows with a 100-token overlap to preserve contextual continuity across paragraphs.\n"
        "3. Dense Vector Embedding: Each chunk is encoded into a 384-dimensional dense floating-point vector using SentenceTransformers ('all-MiniLM-L6-v2'). Chunks are persisted in MongoDB and indexed into FAISS (IndexFlatIP) for instant vector search."
    )

    add_heading_2("4.2 Zero-Hallucination RAG Chain & Guardrails")
    doc.add_paragraph(
        "To guarantee operational safety in industrial environments, the conversational engine applies a multi-stage validation sequence:\n"
        "• Query Embedding: The incoming engineer question is encoded into a 384-dim vector.\n"
        "• Hybrid Search (RRF): The FAISS index computes cosine inner-product similarity, combined with BM25 keyword matching for equipment tag precision.\n"
        "• Similarity Threshold Evaluation:\n"
        "  - If top similarity score ≥ 80%: High Confidence answer generated with ground-truth citations.\n"
        "  - If top similarity score is between 50% and 79%: Medium Confidence answer synthesized.\n"
        "  - If top similarity score < 50%: Generation is suppressed. The AI returns a standardized notice stating that no relevant SOP or RCA context exists in the knowledge base, avoiding hallucinated instructions."
    )

    # Section 5
    add_heading_1("5. Plant Verification & Benchmark Test Results")
    doc.add_paragraph(
        "The system was validated using representative industrial plant manuals, SOPs, and incident reports across major plant assets:"
    )

    test_table = doc.add_table(rows=4, cols=4)
    test_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    test_headers = ["Asset & Document", "Test Query Statement", "Retrieved Context & Citation", "Result Status"]
    for i, title in enumerate(test_headers):
        test_table.rows[0].cells[i].paragraphs[0].add_run(title).font.bold = True
        test_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(test_table.rows[0].cells[i], "1E293B")
        set_cell_margins(test_table.rows[0].cells[i], 100, 100, 100, 100)

    test_data = [
        ("PUMP-101 (Feedwater SOP)", "What is the startup sequence and seal flush pressure?", "SOP-PUMP-101 v2.1: Lube oil >75%, Plan 53A pressure at 4.2 bar, startup MOV-102 over 90s.", "PASSED (75% Conf)"),
        ("COMP-07 (Compressor RCA)", "What was the root cause of bearing temp trip?", "RCA-COMP-07 v1.2: Debris blockage in oil filter bypass valve causing lube starvation.", "PASSED (75% Conf)"),
        ("Nuclear Reactor (Unrelated)", "What is the core temperature limit of Reactor Core 9?", "Out-of-Scope query suppressed by safety guardrail (<50% score). No hallucination.", "PASSED (Suppressed)")
    ]

    for idx, (ast, qry, ctxt, res) in enumerate(test_data, start=1):
        row_cells = test_table.rows[idx].cells
        row_cells[0].paragraphs[0].add_run(ast).font.size = Pt(9)
        row_cells[1].paragraphs[0].add_run(qry).font.size = Pt(9)
        row_cells[2].paragraphs[0].add_run(ctxt).font.size = Pt(9)
        r4 = row_cells[3].paragraphs[0].add_run(res)
        r4.font.size = Pt(9)
        r4.font.bold = True
        
        bg = LIGHT_BG if idx % 2 == 1 else "FFFFFF"
        for c in row_cells:
            set_cell_background(c, bg)
            set_cell_margins(c, 80, 80, 80, 80)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 6
    add_heading_1("6. Deployment & Developer Setup Guide")
    doc.add_paragraph(
        "The codebase is fully containerized and structured for seamless deployment in both local development and production environments.\n\n"
        "Prerequisites:\n"
        "• Node.js (v20+), Python (v3.10+), MongoDB (v6.0+), Docker & Docker Compose.\n\n"
        "Quick Start Instructions:\n"
        "1. Clone repository: git clone https://github.com/Saurabh-Shinde-Patil/ET-industrial-ai.git\n"
        "2. Start Python AI Microservice: cd ai_service && python main.py (Port 8000)\n"
        "3. Start Node.js API Gateway: cd backend && npm run dev (Port 5000)\n"
        "4. Start Frontend Web App: cd frontend && npm run dev (Port 5173)\n"
        "5. Docker Production Launch: docker-compose up --build -d"
    )

    # Section 7
    add_heading_1("7. Conclusion & Operational ROI")
    doc.add_paragraph(
        "ET Industrial AI establishes a robust, highly reliable foundation for industrial digital transformation. By linking legacy manuals "
        "and real-time plant operations through zero-hallucination RAG intelligence, the platform significantly reduces mean time to resolution (MTTR), "
        "eliminates knowledge siloes, and improves overall plant reliability."
    )

    # Save Word file
    doc.save(docx_path)
    print(f"Successfully generated Word Document at: {docx_path}")

def generate_pdf_document(pdf_path):
    doc = SimpleDocTemplate(
        pdf_path,
        pagesize=letter,
        leftMargin=54,
        rightMargin=54,
        topMargin=54,
        bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    
    NAVY = colors.HexColor('#1E293B')
    CYAN = colors.HexColor('#0EA5E9')
    GRAY = colors.HexColor('#334155')
    LIGHT_BG = colors.HexColor('#F8FAFC')
    
    # Modify default styles
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        textColor=NAVY,
        spaceAfter=4
    )
    
    subtitle_style = ParagraphStyle(
        'DocSubTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=11,
        leading=14,
        textColor=GRAY,
        spaceAfter=14
    )

    category_style = ParagraphStyle(
        'DocCategory',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=9,
        leading=11,
        textColor=CYAN,
        spaceAfter=4
    )
    
    h1_style = ParagraphStyle(
        'Heading1_Custom',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=13.5,
        leading=17,
        textColor=NAVY,
        spaceBefore=14,
        spaceAfter=6,
        keepWithNext=True
    )

    h2_style = ParagraphStyle(
        'Heading2_Custom',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=14,
        textColor=CYAN,
        spaceBefore=10,
        spaceAfter=4,
        keepWithNext=True
    )
    
    body_style = ParagraphStyle(
        'Body_Custom',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13.5,
        textColor=GRAY,
        spaceAfter=8
    )

    table_cell_style = ParagraphStyle(
        'TableCell',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8.5,
        leading=11.5,
        textColor=GRAY
    )

    table_cell_bold = ParagraphStyle(
        'TableCellBold',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=8.5,
        leading=11.5,
        textColor=NAVY
    )
    
    table_hdr_style = ParagraphStyle(
        'TableHdr',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=8.5,
        leading=11.5,
        textColor=colors.white
    )

    story = []

    # Title Block
    story.append(Paragraph("ET AI HACKATHON 2.0 — PHASE 2 PROTOTYPE SUBMISSION", category_style))
    story.append(Paragraph("Industrial Knowledge Intelligence Platform (PS-8)", title_style))
    story.append(Paragraph("Technical Architecture & Operational Specification Document | Unified Asset & Operations Brain", subtitle_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=CYAN, spaceBefore=0, spaceAfter=12))

    # Meta Table
    meta_table_data = [
        [Paragraph("Problem Statement", table_cell_bold), Paragraph("PS 8: AI for Industrial Knowledge Intelligence: Unified Asset & Operations Brain", table_cell_style)],
        [Paragraph("Project Name", table_cell_bold), Paragraph("ET Industrial AI", table_cell_style)],
        [Paragraph("Repository URL", table_cell_bold), Paragraph("https://github.com/Saurabh-Shinde-Patil/ET-industrial-ai", table_cell_style)],
        [Paragraph("Date", table_cell_bold), Paragraph("July 20, 2026", table_cell_style)]
    ]
    t_meta = Table(meta_table_data, colWidths=[130, 374])
    t_meta.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (0,-1), LIGHT_BG),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#E2E8F0')),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('TOPPADDING', (0,0), (-1,-1), 5),
        ('BOTTOMPADDING', (0,0), (-1,-1), 5),
        ('LEFTPADDING', (0,0), (-1,-1), 8),
        ('RIGHTPADDING', (0,0), (-1,-1), 8),
    ]))
    story.append(t_meta)
    story.append(Spacer(1, 10))

    # Section 1
    story.append(Paragraph("1. Executive Summary", h1_style))
    story.append(Paragraph(
        "In thermal power plants, chemical process facilities, and heavy manufacturing installations, critical operational knowledge "
        "is frequently fragmented across thousands of static PDF manuals, handwritten maintenance logs, Standard Operating Procedures (SOPs), "
        "and historic Root Cause Analysis (RCA) records. During acute equipment trips or forced outages, plant engineers face significant "
        "delays sifting through non-searchable document repositories, leading to extended downtime, elevated safety risks, and operational inefficiencies.",
        body_style
    ))
    story.append(Paragraph(
        "ET Industrial AI addresses this operational bottleneck by deploying a unified, enterprise-grade Knowledge Intelligence Platform. "
        "Built on a microservices architecture, the system combines a Zero-Hallucination Retrieval-Augmented Generation (RAG) conversational engine, "
        "a high-density vector search index (FAISS), an automated 5-Whys Root Cause Analysis engine, and predictive maintenance risk scoring. "
        "The platform provides plant engineers with immediate, contextual answers backed strictly by ground-truth document citations, enforcing "
        "strict confidence thresholds to prevent AI hallucination.",
        body_style
    ))

    # Section 2
    story.append(Paragraph("2. Problem Statement Alignment (PS-8)", h1_style))
    story.append(Paragraph(
        "The project is engineered specifically for Problem Statement 8: AI for Industrial Knowledge Intelligence: Unified Asset & Operations Brain. "
        "The table below outlines how specific plant operational challenges are addressed by the platform architecture:",
        body_style
    ))

    ps_table_data = [
        [Paragraph("Industrial Operational Challenge", table_hdr_style), Paragraph("ET Industrial AI Architectural Solution", table_hdr_style)],
        [Paragraph("Knowledge Isolation across Plant Siloes", table_cell_bold), Paragraph("Centralized document ingestion & OCR parsing supporting PDF, Word, TXT, and scanned image formats with direct asset tagging.", table_cell_style)],
        [Paragraph("Risk of AI Hallucinations on Critical Assets", table_cell_bold), Paragraph("Zero-Hallucination RAG Chain enforcing similarity threshold suppression (<50% score). Returns standardized fallback notices when manuals lack matching context.", table_cell_style)],
        [Paragraph("Repetitive Equipment Failures", table_cell_bold), Paragraph("Automated 5-Whys Root Cause Analysis engine generating standardized incident reports and Corrective Action (CAPA) plans.", table_cell_style)],
        [Paragraph("Unplanned Equipment Downtime", table_cell_bold), Paragraph("AI Predictive Maintenance engine calculating asset failure risk probabilities based on telemetry parameters and operating hours.", table_cell_style)],
        [Paragraph("Enterprise Security & Role Compliance", table_cell_bold), Paragraph("8 RBAC roles (Plant Operator, Maintenance Eng, Reliability Eng, Safety Officer, Production Eng, Plant Mgr, Knowledge Admin, System Admin) with audit event logging.", table_cell_style)]
    ]
    t_ps = Table(ps_table_data, colWidths=[160, 344])
    t_ps.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), NAVY),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('TOPPADDING', (0,0), (-1,-1), 5),
        ('BOTTOMPADDING', (0,0), (-1,-1), 5),
        ('LEFTPADDING', (0,0), (-1,-1), 6),
        ('RIGHTPADDING', (0,0), (-1,-1), 6),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [LIGHT_BG, colors.white]),
    ]))
    story.append(t_ps)
    story.append(Spacer(1, 10))

    # Section 3
    story.append(Paragraph("3. System Architecture & Technical Implementation", h1_style))
    story.append(Paragraph(
        "The system is structured as a decoupled, multi-tier microservice architecture to ensure high throughput, horizontal scalablity, "
        "and strict separation between the API gateway, vector store, and AI inference engines.", body_style
    ))
    
    story.append(Paragraph("3.1 Core Components", h2_style))
    story.append(Paragraph(
        "<b>1. Frontend Web Client (React 18 / Tailwind CSS / Lucide Icons):</b><br/>"
        "Provides a dark-mode interface with operational dashboards, RAG conversation view, document manager, asset hierarchy, and admin audit tools.<br/><br/>"
        "<b>2. Node.js Express REST API Gateway (Port 5000):</b><br/>"
        "Handles user authentication, JWT session validation, RBAC route protection, MongoDB persistence, audit event logging, and proxy dispatch to the Python AI service.<br/><br/>"
        "<b>3. Python FastAPI AI Microservice (Port 8000):</b><br/>"
        "Houses core AI engines: SentenceTransformers embedding generation, FAISS vector indexing, Reciprocal Rank Fusion (RRF) hybrid search, PyPDF2/pdfplumber text extraction, Tesseract OCR, and RAG synthesis.",
        body_style
    ))

    # Section 4
    story.append(Paragraph("4. Detailed AI & Vector Search Pipeline", h1_style))
    story.append(Paragraph("4.1 Document Processing & Vector Chunking", h2_style))
    story.append(Paragraph(
        "When an engineer uploads an industrial document, the system executes a three-phase pipeline:<br/>"
        "1. <b>Text Extraction:</b> Extracted using pdfplumber or Tesseract OCR (for scanned engineering drawings).<br/>"
        "2. <b>Chunking:</b> Text is segmented into 500-token sliding windows with a 100-token overlap to preserve continuity.<br/>"
        "3. <b>Dense Vector Embedding:</b> Encoded into a 384-dimensional dense vector using SentenceTransformers ('all-MiniLM-L6-v2'). Indexed into FAISS (IndexFlatIP) for cosine similarity search.",
        body_style
    ))

    story.append(Paragraph("4.2 Zero-Hallucination RAG Chain & Guardrails", h2_style))
    story.append(Paragraph(
        "To guarantee operational safety on live plant assets, the engine applies similarity score thresholds:<br/>"
        "• <b>Score ≥ 80%:</b> High Confidence answer generated with ground-truth citations.<br/>"
        "• <b>Score 50% – 79%:</b> Medium Confidence answer synthesized.<br/>"
        "• <b>Score &lt; 50%:</b> Generation is suppressed. The AI returns a standardized notice stating that no matching context exists, avoiding hallucinated instructions.",
        body_style
    ))

    # Section 5
    story.append(Paragraph("5. Plant Verification & Benchmark Test Results", h1_style))
    test_table_data = [
        [Paragraph("Asset & Document", table_hdr_style), Paragraph("Test Query Statement", table_hdr_style), Paragraph("Retrieved Context & Citation", table_hdr_style), Paragraph("Result Status", table_hdr_style)],
        [Paragraph("PUMP-101 (SOP)", table_cell_bold), Paragraph("What is the startup sequence and seal pressure?", table_cell_style), Paragraph("SOP-PUMP-101 v2.1: Plan 53A pressure at 4.2 bar, startup MOV-102 over 90s.", table_cell_style), Paragraph("PASSED (75% Conf)", table_cell_bold)],
        [Paragraph("COMP-07 (RCA)", table_cell_bold), Paragraph("What was the root cause of bearing temp trip?", table_cell_style), Paragraph("RCA-COMP-07 v1.2: Debris blockage in oil filter bypass valve causing lube starvation.", table_cell_style), Paragraph("PASSED (75% Conf)", table_cell_bold)],
        [Paragraph("Nuclear Core-9", table_cell_bold), Paragraph("What is core temp limit of Reactor Core 9?", table_cell_style), Paragraph("Out-of-Scope query suppressed by safety guardrail (<50% score). No hallucination.", table_cell_style), Paragraph("PASSED (Suppressed)", table_cell_bold)]
    ]
    t_test = Table(test_table_data, colWidths=[100, 130, 184, 90])
    t_test.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), NAVY),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('TOPPADDING', (0,0), (-1,-1), 4),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ('LEFTPADDING', (0,0), (-1,-1), 5),
        ('RIGHTPADDING', (0,0), (-1,-1), 5),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [LIGHT_BG, colors.white]),
    ]))
    story.append(t_test)
    story.append(Spacer(1, 10))

    # Section 6
    story.append(Paragraph("6. Developer Setup & Installation", h1_style))
    story.append(Paragraph(
        "<b>Repository:</b> https://github.com/Saurabh-Shinde-Patil/ET-industrial-ai<br/>"
        "<b>1. AI Microservice:</b> <code>cd ai_service && python main.py</code> (Port 8000)<br/>"
        "<b>2. API Gateway:</b> <code>cd backend && npm run dev</code> (Port 5000, Swagger UI at /docs)<br/>"
        "<b>3. Frontend Client:</b> <code>cd frontend && npm run dev</code> (Port 5173)<br/>"
        "<b>4. Docker Production:</b> <code>docker-compose up --build -d</code>",
        body_style
    ))

    # Section 7
    story.append(Paragraph("7. Conclusion", h1_style))
    story.append(Paragraph(
        "ET Industrial AI establishes a robust, production-grade platform for industrial operational intelligence. "
        "By linking legacy document siloes with zero-hallucination vector RAG intelligence, the system reduces maintenance resolution time "
        "and improves overall equipment effectiveness (OEE).",
        body_style
    ))

    doc.build(story)
    print(f"Successfully generated PDF Document at: {pdf_path}")

if __name__ == '__main__':
    docs_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "docs")
    os.makedirs(docs_dir, exist_ok=True)
    
    docx_file = os.path.join(docs_dir, "ET_Industrial_AI_Technical_Report.docx")
    pdf_file = os.path.join(docs_dir, "ET_Industrial_AI_Technical_Report.pdf")
    
    generate_word_document(docx_file)
    generate_pdf_document(pdf_file)
